"""
generate_samples.py — Create synthetic D&D sample documents for the RAG knowledge base.

Generates 10–20 original items per category using Claude, then saves them
as .docx files in data/samples/<Category>/.  These files can be ingested
like any other source document and are safe to share in the public repo.

Usage:
    python src/generate_samples.py                      # all categories
    python src/generate_samples.py --category weapon    # one category
    python src/generate_samples.py --count 15           # 15 per category
"""

import argparse
import json
import sys
import time
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

sys.path.insert(0, str(Path(__file__).resolve().parent))
from generate import generate

PROJECT_ROOT = Path(__file__).resolve().parent.parent
# Samples go directly into the category data dirs so ingest.py picks them up automatically.
# Files are prefixed Sample_ so contributors can distinguish them from their own source docs.
SAMPLES_DIR  = PROJECT_ROOT / "data"

# Varied seed parameters so the LLM produces diverse items
SEEDS = {
    "weapon": [
        {"rarity": "common",    "type": "dagger",    "theme": "shadow",   "location": "Waterdeep"},
        {"rarity": "uncommon",  "type": "longsword", "theme": "frost",    "location": "Icewind Dale"},
        {"rarity": "rare",      "type": "shortbow",  "theme": "fire",     "location": "Calimshan"},
        {"rarity": "very rare", "type": "staff",     "theme": "arcane",   "location": "Candlekeep"},
        {"rarity": "legendary", "type": "greatsword","theme": "lightning","location": "Baldur's Gate"},
        {"rarity": "uncommon",  "type": "rapier",    "theme": "poison",   "location": "Amn"},
        {"rarity": "rare",      "type": "warhammer", "theme": "divine",   "location": "Silverymoon"},
        {"rarity": "common",    "type": "handaxe",   "theme": "nature",   "location": "High Forest"},
        {"rarity": "very rare", "type": "crossbow",  "theme": "necrotic", "location": "Undermountain"},
        {"rarity": "legendary", "type": "trident",   "theme": "storm",    "location": "Neverwinter"},
    ],
    "npc": [
        {"char_class": "Rogue",   "rarity": "uncommon", "theme": "spy",        "location": "Waterdeep"},
        {"char_class": "Wizard",  "rarity": "rare",     "theme": "scholar",    "location": "Candlekeep"},
        {"char_class": "Paladin", "rarity": "uncommon", "theme": "zealot",     "location": "Baldur's Gate"},
        {"char_class": "Ranger",  "rarity": "common",   "theme": "tracker",    "location": "Neverwinter Wood"},
        {"char_class": "Warlock", "rarity": "rare",     "theme": "seeker",     "location": "Luskan"},
        {"char_class": "Cleric",  "rarity": "uncommon", "theme": "healer",     "location": "Silverymoon"},
        {"char_class": "Fighter", "rarity": "common",   "theme": "mercenary",  "location": "Amn"},
        {"char_class": "Bard",    "rarity": "uncommon", "theme": "spy",        "location": "Calimport"},
        {"char_class": "Druid",   "rarity": "rare",     "theme": "guardian",   "location": "High Forest"},
        {"char_class": "Sorcerer","rarity": "legendary","theme": "wild magic", "location": "Anauroch"},
    ],
    "monster": [
        {"cr": "1",  "type": "beast",     "theme": "forest",    "location": "Neverwinter Wood"},
        {"cr": "3",  "type": "undead",    "theme": "cold",      "location": "Icewind Dale"},
        {"cr": "5",  "type": "humanoid",  "theme": "tribal",    "location": "Sword Mountains"},
        {"cr": "8",  "type": "fiend",     "theme": "shadow",    "location": "Undermountain"},
        {"cr": "10", "type": "dragon",    "theme": "fire",      "location": "Calimshan"},
        {"cr": "2",  "type": "fey",       "theme": "trickster", "location": "High Forest"},
        {"cr": "6",  "type": "construct", "theme": "arcane",    "location": "Candlekeep"},
        {"cr": "12", "type": "aberration","theme": "psionic",   "location": "Underdark"},
        {"cr": "4",  "type": "undead",    "theme": "cursed",    "location": "Baldur's Gate sewers"},
        {"cr": "7",  "type": "giant",     "theme": "storm",     "location": "North"},
    ],
    "artifact": [
        {"rarity": "legendary", "type": "ring",    "theme": "shadow",  "location": "Waterdeep"},
        {"rarity": "artifact",  "type": "crown",   "theme": "divine",  "location": "Silverymoon"},
        {"rarity": "very rare", "type": "amulet",  "theme": "arcane",  "location": "Candlekeep"},
        {"rarity": "legendary", "type": "staff",   "theme": "storm",   "location": "Sea of Swords"},
        {"rarity": "artifact",  "type": "tome",    "theme": "necrotic","location": "Undermountain"},
        {"rarity": "very rare", "type": "cloak",   "theme": "illusion","location": "Amn"},
        {"rarity": "legendary", "type": "gauntlets","theme":"strength","location": "Luskan"},
        {"rarity": "artifact",  "type": "mirror",  "theme": "divination","location": "Anauroch"},
        {"rarity": "very rare", "type": "orb",     "theme": "fire",    "location": "Calimshan"},
        {"rarity": "legendary", "type": "sword",   "theme": "holy",    "location": "Baldur's Gate"},
    ],
    "location": [
        {"type": "dungeon",  "terrain": "underground", "theme": "undead",   "rarity": "Landmark"},
        {"type": "city",     "terrain": "coastal",     "theme": "political","rarity": "City"},
        {"type": "fortress", "terrain": "mountain",    "theme": "military", "rarity": "Landmark"},
        {"type": "ruins",    "terrain": "forest",      "theme": "ancient",  "rarity": "Landmark"},
        {"type": "tower",    "terrain": "plains",      "theme": "arcane",   "rarity": "Landmark"},
        {"type": "town",     "terrain": "coastal",     "theme": "trade",    "rarity": "Town"},
        {"type": "temple",   "terrain": "mountain",    "theme": "divine",   "rarity": "Landmark"},
        {"type": "cavern",   "terrain": "underground", "theme": "fey",      "rarity": "Landmark"},
        {"type": "village",  "terrain": "forest",      "theme": "frontier", "rarity": "Village"},
        {"type": "port",     "terrain": "coastal",     "theme": "pirates",  "rarity": "Town"},
    ],
}

# Map category id → folder name (matches ingest.py CATEGORY_DIRS)
FOLDER_MAP = {
    "weapon":   "Weapons",
    "npc":      "Characters",
    "monster":  "Monsters",
    "artifact": "Artifacts",
    "location": "Locations",
}


def _write_docx(items: list[dict], category: str, path: Path) -> None:
    doc = Document()

    # Title
    title = doc.add_heading(f"Synthetic {category.title()} Samples", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x6a, 0x1d, 0x15)

    doc.add_paragraph(
        "These items were generated by the Artificer's Codex LLM pipeline "
        "and are original synthetic content safe to distribute."
    )
    doc.add_paragraph("")

    for i, item in enumerate(items, 1):
        name = item.get("name", f"Item {i}")
        heading = doc.add_heading(f"{i}. {name}", level=2)
        heading.runs[0].font.color.rgb = RGBColor(0x3a, 0x19, 0x0b)

        for key, value in item.items():
            if key in ("name", "source_category", "image_prompt"):
                continue
            label = key.replace("_", " ").title()
            para = doc.add_paragraph()
            run_label = para.add_run(f"{label}: ")
            run_label.bold = True
            if isinstance(value, list):
                if value and isinstance(value[0], dict):
                    text = "; ".join(
                        f"{v.get('name','')}: {v.get('description','')}"
                        for v in value
                    )
                else:
                    text = "; ".join(str(v) for v in value)
            else:
                text = str(value) if value is not None else "—"
            para.add_run(text)

        doc.add_paragraph("")

    doc.save(str(path))
    print(f"  Saved {len(items)} items → {path.name}")


def generate_category(category: str, count: int) -> None:
    seeds = SEEDS.get(category, [{}])
    folder = SAMPLES_DIR / FOLDER_MAP[category]
    folder.mkdir(parents=True, exist_ok=True)

    items = []
    total = min(count, len(seeds))
    print(f"\n{category.upper()} — generating {total} samples…")

    for i, seed in enumerate(seeds[:total]):
        print(f"  [{i+1}/{total}] {seed}…", end=" ", flush=True)
        try:
            content = generate(category=category, **seed)
            items.append(content)
            print(f"✓ {content.get('name','?')}")
        except Exception as e:
            print(f"✗ {e}")
        if i < total - 1:
            time.sleep(1)  # rate-limit courtesy pause

    if items:
        batch_num = len(list(folder.glob("Sample_*.docx"))) + 1
        out_path = folder / f"Sample_{FOLDER_MAP[category]}_Batch{batch_num}.docx"
        _write_docx(items, category, out_path)


def main():
    parser = argparse.ArgumentParser(description="Generate synthetic RAG sample documents")
    parser.add_argument("--category", choices=list(SEEDS.keys()), help="Single category to generate")
    parser.add_argument("--count", type=int, default=10, help="Items per category (default 10)")
    args = parser.parse_args()

    categories = [args.category] if args.category else list(SEEDS.keys())

    print(f"Generating {args.count} synthetic samples for: {', '.join(categories)}")
    print("These will be saved to data/samples/ and are safe to commit.\n")

    for cat in categories:
        generate_category(cat, args.count)

    print("\nDone. Run 'python src/ingest.py' to add these to ChromaDB.")


if __name__ == "__main__":
    main()
