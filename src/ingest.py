"""
ingest.py — Parse .docx files across all D&D categories and index them
into ChromaDB via LlamaIndex.

Usage:
    python src/ingest.py                    # Ingest all categories
    python src/ingest.py --category weapon  # Ingest one category
"""

import argparse
import re
import sys
from pathlib import Path

import chromadb
from docx import Document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from dotenv import load_dotenv
from llama_index.core import Settings, StorageContext, VectorStoreIndex
from llama_index.core import Document as LlamaDocument
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore
from rich.console import Console
from rich.table import Table

load_dotenv()

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = PROJECT_ROOT / "data"
CHROMA_PATH = PROJECT_ROOT / "chroma_db"
COLLECTION_NAME = "dnd_knowledge"

console = Console()

# Category → data directory name mapping
CATEGORY_DIRS = {
    "weapon": "Weapons",
    "npc": "Characters",
    "location": "Locations",
    "monster": "Monsters",
    "artifact": "Artifacts",
}

# Chunking strategy per category (from project instructions)
CHUNK_CONFIG = {
    "weapon":   {"chunk_size": 256, "overlap": 50},
    "npc":      {"chunk_size": 256, "overlap": 50},
    "artifact": {"chunk_size": 256, "overlap": 50},
    "location": {"chunk_size": 256, "overlap": 50},
    "monster":  {"chunk_size": 256, "overlap": 50},
}


# ---------------------------------------------------------------------------
# Shared docx helpers
# ---------------------------------------------------------------------------

def _iter_body(doc: Document):
    """Yield paragraphs and tables in their actual document order."""
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield DocxParagraph(child, body)
        elif tag == "tbl":
            yield DocxTable(child, body)


def _para_style_id(paragraph: DocxParagraph) -> str:
    pPr = paragraph._element.pPr
    if pPr is None:
        return ""
    pStyle = pPr.pStyle
    return pStyle.val if pStyle is not None else ""


def _table_to_dict(table: DocxTable) -> dict:
    """Convert a 2-column key/value table into a dict."""
    data = {}
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2 and cells[0]:
            data[cells[0]] = cells[1]
    return data


def _table_cell_text(table: DocxTable) -> str:
    """Get all text from a table (for single-cell section header tables)."""
    return "\n".join(c.text.strip() for row in table.rows for c in row.cells if c.text.strip())


# ---------------------------------------------------------------------------
# Weapon parser
# ---------------------------------------------------------------------------

_NAME_RE = re.compile(r"^\d+\.\s+(.+)$")
_WEAPON_FIELD_KEYS = {"Item Type", "Rarity", "Requires Attunement"}

_WEAPON_FIELD_MAP = {
    "Item Type":            "item_type",
    "Rarity":               "rarity",
    "Requires Attunement":  "requires_attunement",
    "Attunement Details":   "attunement_details",
    "Physical Description": "physical_description",
    "Properties":           "properties",
    "Special Ability":      "special_ability",
    "Spells":               "spells",
    "Lore & History":       "lore_history",
    "Special Conditions":   "special_conditions",
    "Curse":                "curse",
    "Sentience":            "sentience",
    "Image Prompt":         "image_prompt",
}


def _extract_weapon_name(paragraph: DocxParagraph) -> str | None:
    text = paragraph.text.strip()
    if not text:
        return None
    m = _NAME_RE.match(text)
    if m:
        return m.group(1).strip()
    style_id = _para_style_id(paragraph)
    if style_id in ("2", "Heading2", "heading2"):
        return text
    return None


def _is_template_table(table: DocxTable) -> bool:
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2:
            continue
        key, value = cells[0], cells[1]
        if key == "Rarity" and "/" in value:
            return True
        if key == "Name" and len(value) > 40:
            return True
    return False


def _is_weapon_table(table: DocxTable) -> bool:
    if _is_template_table(table):
        return False
    first_col = {row.cells[0].text.strip() for row in table.rows if row.cells}
    return bool(first_col & _WEAPON_FIELD_KEYS)


def _parse_weapon_table(table: DocxTable) -> dict:
    data = {}
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2:
            continue
        key, value = cells[0], cells[1]
        field = _WEAPON_FIELD_MAP.get(key)
        if field is None:
            continue
        if field == "special_ability" and "special_ability" in data:
            data[field] = data[field] + "\n" + value
        else:
            data[field] = value
    return data


def extract_weapons(docx_path: Path) -> list[dict]:
    doc = Document(docx_path)
    weapons = []
    current_name = None

    for element in _iter_body(doc):
        if isinstance(element, DocxParagraph):
            name = _extract_weapon_name(element)
            if name:
                current_name = name
        elif isinstance(element, DocxTable):
            if not _is_weapon_table(element):
                continue
            fields = _parse_weapon_table(element)
            if not fields:
                continue
            name = current_name or fields.pop("name", None) or "Unknown"
            current_name = None
            weapons.append({"name": name, **fields})

    return weapons


def _weapon_to_text(weapon: dict) -> str:
    lines = [f"Name: {weapon.get('name', 'Unknown')}"]
    for field, key in _WEAPON_FIELD_MAP.items():
        value = weapon.get(key, "")
        if value and value.upper() != "N/A":
            lines.append(f"{field}: {value}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# NPC parser
# ---------------------------------------------------------------------------

# Section labels from 1×1 header tables
_NPC_SECTIONS = {"NARRATIVE", "SOCIAL", "GAMEPLAY", "ADVENTURE", "RAG / KNOWLEDGE BASE"}

# Field labels within section paragraphs (prefix before first ": ")
_NPC_NARRATIVE_FIELDS = [
    "Appearance", "Personality", "Background", "Motivations & Goals", "Secrets"
]
_NPC_SOCIAL_FIELDS = ["Faction / Affiliation", "Relationships"]
_NPC_GAMEPLAY_FIELDS = ["Notable Abilities", "Equipment", "Combat Role", "Roleplaying Notes"]
_NPC_ADVENTURE_FIELDS = ["Encounter Hooks", "Rumors"]
_NPC_RAG_FIELDS = ["Image Prompt", "Tags", "Source Books"]


def extract_npcs(docx_path: Path) -> list[dict]:
    doc = Document(docx_path)
    npcs = []
    elements = list(_iter_body(doc))

    i = 0
    while i < len(elements):
        el = elements[i]

        # Look for name-block tables (1 row, 1 col, with NPC name)
        if isinstance(el, DocxTable) and len(el.rows) == 1 and len(el.columns) <= 1:
            cell_text = el.rows[0].cells[0].text.strip()
            lines = [l.strip() for l in cell_text.split("\n") if l.strip()]

            # Name block has the name on line 0 and "Race · Gender · Class · Alignment" on line 1
            if len(lines) >= 2 and "·" in lines[1]:
                npc_name = lines[0]
                i += 1

                # Next should be the 8-row metadata table
                if i < len(elements) and isinstance(elements[i], DocxTable):
                    meta = _table_to_dict(elements[i])
                    i += 1
                else:
                    meta = {}

                # Collect section content
                sections = {}
                current_section = None

                while i < len(elements):
                    el2 = elements[i]

                    if isinstance(el2, DocxTable):
                        header_text = _table_cell_text(el2).strip()

                        # Check if it's a section header
                        if header_text in _NPC_SECTIONS:
                            current_section = header_text
                            sections[current_section] = []
                            i += 1
                            continue

                        # Check if it's the next NPC's name block
                        if len(el2.rows) == 1 and len(el2.columns) <= 1:
                            cell = el2.rows[0].cells[0].text.strip()
                            clines = [l.strip() for l in cell.split("\n") if l.strip()]
                            if len(clines) >= 2 and "·" in clines[1]:
                                break  # Next NPC

                        # Multi-row table within section — treat as metadata
                        if len(el2.rows) > 1 and current_section:
                            sections[current_section].append(
                                _table_cell_text(el2)
                            )
                        i += 1
                        continue

                    if isinstance(el2, DocxParagraph):
                        text = el2.text.strip()
                        if text and current_section:
                            sections[current_section].append(text)
                        i += 1
                        continue

                    i += 1

                # Build the NPC dict from collected sections
                npc = _build_npc_dict(npc_name, meta, sections)
                npcs.append(npc)
                continue

        i += 1

    return npcs


def _extract_section_fields(paragraphs: list[str], field_names: list[str]) -> dict:
    """Parse 'Label: value' paragraphs into a dict."""
    result = {}
    current_key = None
    current_values = []

    for para in paragraphs:
        # Check if this paragraph starts with a known field label
        matched = False
        for field in field_names:
            if para.startswith(field + ":"):
                # Save previous field
                if current_key:
                    result[current_key] = "\n".join(current_values).strip()
                current_key = field
                current_values = [para[len(field) + 1:].strip()]
                matched = True
                break

        if not matched and current_key:
            current_values.append(para)

    if current_key:
        result[current_key] = "\n".join(current_values).strip()

    return result


def _build_npc_dict(name: str, meta: dict, sections: dict) -> dict:
    narrative = _extract_section_fields(
        sections.get("NARRATIVE", []), _NPC_NARRATIVE_FIELDS
    )
    social = _extract_section_fields(
        sections.get("SOCIAL", []), _NPC_SOCIAL_FIELDS
    )
    gameplay = _extract_section_fields(
        sections.get("GAMEPLAY", []), _NPC_GAMEPLAY_FIELDS
    )
    adventure = _extract_section_fields(
        sections.get("ADVENTURE", []), _NPC_ADVENTURE_FIELDS
    )
    rag = _extract_section_fields(
        sections.get("RAG / KNOWLEDGE BASE", []), _NPC_RAG_FIELDS
    )

    # Parse encounter hooks and rumors into lists
    hooks_text = adventure.get("Encounter Hooks", "")
    hooks = [h.strip() for h in hooks_text.split("\n") if h.strip()] if hooks_text else []

    rumors_text = adventure.get("Rumors", "")
    rumors = [r.strip() for r in rumors_text.split("\n") if r.strip()] if rumors_text else []

    return {
        "name": name,
        "archetype": meta.get("Archetype", ""),
        "race": meta.get("Race", ""),
        "gender": meta.get("Gender", ""),
        "char_class": meta.get("Class / Level", ""),
        "challenge_rating": meta.get("Challenge Rating", ""),
        "alignment": meta.get("Alignment", ""),
        "aliases": meta.get("Aliases", ""),
        "region": meta.get("Region", ""),
        "appearance": narrative.get("Appearance", ""),
        "personality": narrative.get("Personality", ""),
        "backstory": narrative.get("Background", ""),
        "motivations": narrative.get("Motivations & Goals", ""),
        "secrets": narrative.get("Secrets", ""),
        "affiliation": social.get("Faction / Affiliation", ""),
        "relationships": social.get("Relationships", ""),
        "abilities_and_skills": gameplay.get("Notable Abilities", ""),
        "equipment": gameplay.get("Equipment", ""),
        "combat_tactics": gameplay.get("Combat Role", ""),
        "roleplaying_tips": gameplay.get("Roleplaying Notes", ""),
        "encounter_hooks": hooks,
        "rumors": rumors,
        "image_prompt": rag.get("Image Prompt", ""),
        "tags": rag.get("Tags", ""),
    }


def _npc_to_text(npc: dict) -> str:
    lines = [f"Name: {npc['name']}"]
    skip = {"name", "encounter_hooks", "rumors", "image_prompt", "tags"}
    for key, value in npc.items():
        if key in skip or not value:
            continue
        label = key.replace("_", " ").title()
        lines.append(f"{label}: {value}")
    if npc.get("encounter_hooks"):
        lines.append("Encounter Hooks: " + " | ".join(npc["encounter_hooks"]))
    if npc.get("rumors"):
        lines.append("Rumors: " + " | ".join(npc["rumors"]))
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Location parser
# ---------------------------------------------------------------------------

# Section headers are all-caps paragraphs within a location entry
_LOCATION_SECTIONS = {
    "FACTIONS & POWER GROUPS", "NOTABLE NPCS", "NOTABLE LANDMARKS",
    "HISTORY", "LORE & ATMOSPHERE", "ECONOMY & TRADE", "TERRAIN & CLIMATE",
    "SECRETS (DM ONLY)", "ADVENTURE HOOKS", "TYPICAL ENCOUNTERS",
    "DANGERS & THREATS", "RUMORS & HEARSAY", "IMAGE PROMPT (DALL-E 3)",
    # Variations for different batches
    "FACTIONS", "NPCS", "LANDMARKS", "SECRETS", "HOOKS", "ENCOUNTERS",
    "DANGERS", "RUMORS", "IMAGE PROMPT",
}


def _is_location_header(text: str) -> bool:
    """Check if a paragraph is a location name header (all-caps, short)."""
    if not text or len(text) < 3:
        return False
    # Must be mostly uppercase letters/spaces, not a section header
    cleaned = text.replace("'", "").replace("'", "")
    if not cleaned.replace(" ", "").replace("-", "").isalpha():
        return False
    if text != text.upper():
        return False
    if text in _LOCATION_SECTIONS:
        return False
    return True


def extract_locations(docx_path: Path) -> list[dict]:
    doc = Document(docx_path)
    elements = list(_iter_body(doc))
    locations = []
    i = 0

    while i < len(elements):
        el = elements[i]

        # Look for location name headers (all-caps paragraphs)
        if isinstance(el, DocxParagraph):
            text = el.text.strip()

            if _is_location_header(text):
                loc_name = text.title()  # Convert "WATERDEEP" → "Waterdeep"
                i += 1

                # Collect epithets line (has "|" separators)
                epithets = ""
                if i < len(elements) and isinstance(elements[i], DocxParagraph):
                    next_text = elements[i].text.strip()
                    if "|" in next_text:
                        epithets = next_text
                        i += 1

                # Collect type/region line
                loc_type = ""
                region = ""
                if i < len(elements) and isinstance(elements[i], DocxParagraph):
                    next_text = elements[i].text.strip()
                    # Pattern: "CITY     Sword Coast (Central)     Northern Sword Coast..."
                    parts = [p.strip() for p in re.split(r"\s{3,}", next_text) if p.strip()]
                    if parts:
                        loc_type = parts[0]
                        region = ", ".join(parts[1:]) if len(parts) > 1 else ""
                        i += 1

                # Look for stats table (Population, Government, etc.)
                stats = {}
                if i < len(elements) and isinstance(elements[i], DocxTable):
                    stats = _table_to_dict(elements[i])
                    i += 1

                # Collect section content until next location header
                sections = {}
                current_section = "description"
                sections[current_section] = []

                while i < len(elements):
                    el2 = elements[i]

                    if isinstance(el2, DocxParagraph):
                        t = el2.text.strip()
                        if not t:
                            i += 1
                            continue

                        # Check for next location header
                        if _is_location_header(t):
                            break

                        # Check for section header
                        if t.upper() == t and t in _LOCATION_SECTIONS:
                            current_section = t.lower()
                            sections.setdefault(current_section, [])
                            i += 1
                            continue

                        sections.setdefault(current_section, []).append(t)

                    elif isinstance(el2, DocxTable):
                        # Stats tables for sub-locations, skip
                        pass

                    i += 1

                # Build location dict
                loc = _build_location_dict(
                    loc_name, loc_type, region, epithets, stats, sections
                )
                locations.append(loc)
                continue

        i += 1

    return locations


def _build_location_dict(
    name: str, loc_type: str, region: str, epithets: str,
    stats: dict, sections: dict
) -> dict:
    # Combine description from various sections
    desc_parts = sections.get("description", [])
    history_parts = sections.get("history", [])
    lore_parts = sections.get("lore & atmosphere", [])

    # Get image prompt from the last section
    img_sections = (
        sections.get("image prompt (dall-e 3)", [])
        or sections.get("image prompt", [])
    )
    image_prompt = img_sections[0] if img_sections else ""
    tags = img_sections[1] if len(img_sections) > 1 else ""

    # Collect hooks
    hook_parts = (
        sections.get("adventure hooks", [])
        or sections.get("hooks", [])
    )

    # Collect rumors
    rumor_parts = (
        sections.get("rumors & hearsay", [])
        or sections.get("rumors", [])
    )

    # NPCs section
    npc_parts = (
        sections.get("notable npcs", [])
        or sections.get("npcs", [])
    )

    # Landmarks / notable features
    feature_parts = (
        sections.get("notable landmarks", [])
        or sections.get("landmarks", [])
    )

    return {
        "name": name,
        "location_type": loc_type,
        "region": region,
        "epithets": epithets,
        "population": stats.get("Population", ""),
        "government": stats.get("Government", ""),
        "alignment": stats.get("Alignment", ""),
        "factions": stats.get("Factions", ""),
        "description": "\n".join(desc_parts + lore_parts),
        "notable_features": "\n".join(feature_parts),
        "history": "\n".join(history_parts),
        "npcs": "\n".join(npc_parts),
        "hooks": "\n".join(hook_parts),
        "rumors": "\n".join(rumor_parts),
        "image_prompt": image_prompt,
        "tags": tags,
    }


def _location_to_text(loc: dict) -> str:
    lines = [f"Name: {loc['name']}"]
    skip = {"name", "image_prompt", "tags"}
    for key, value in loc.items():
        if key in skip or not value:
            continue
        label = key.replace("_", " ").title()
        lines.append(f"{label}: {value}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Monster parser
# ---------------------------------------------------------------------------

_MONSTER_STAT_KEYS = {
    "Monster Type", "Challenge Rating", "Armor Class", "Hit Points",
    "Speed", "Ability Scores", "Saving Throws", "Skills",
    "Damage Resistances", "Damage Immunities", "Condition Immunities",
    "Senses", "Languages", "Dungeon Type", "Faction",
}

_MONSTER_SECTION_HEADERS = {
    "TRAITS", "ACTIONS", "REACTIONS", "LORE", "ECOLOGY",
    "ENCOUNTER HOOKS", "IMAGE PROMPT",
}


def _is_monster_stat_table(table: DocxTable) -> bool:
    first_col = {row.cells[0].text.strip() for row in table.rows if row.cells}
    return bool(first_col & {"Monster Type", "Challenge Rating", "Armor Class"})


def extract_monsters(docx_path: Path) -> list[dict]:
    doc = Document(docx_path)
    elements = list(_iter_body(doc))
    monsters = []
    i = 0

    while i < len(elements):
        el = elements[i]

        # Monster entries start with a stat block table
        if isinstance(el, DocxTable) and _is_monster_stat_table(el):
            stats = _table_to_dict(el)
            i += 1

            # The name paragraph precedes the table — look back
            # Actually, based on the data: name paragraph comes before the table
            # We need to have captured it. Let's look at the preceding paragraphs.
            monster_name = "Unknown"

            # Walk back to find the name (non-empty, non-section-header paragraph)
            for j in range(i - 2, max(i - 5, -1), -1):
                if j >= 0 and isinstance(elements[j], DocxParagraph):
                    t = elements[j].text.strip()
                    if t and t not in _MONSTER_SECTION_HEADERS and t.upper() != t:
                        monster_name = t
                        break
                    # Also check if it's a title-cased name in all caps section
                    if t and t not in _MONSTER_SECTION_HEADERS:
                        monster_name = t
                        break

            # Collect section paragraphs until next stat table
            sections = {}
            current_section = "traits"  # Default section after stat block

            while i < len(elements):
                el2 = elements[i]

                if isinstance(el2, DocxTable):
                    if _is_monster_stat_table(el2):
                        break  # Next monster
                    i += 1
                    continue

                if isinstance(el2, DocxParagraph):
                    t = el2.text.strip()
                    if not t:
                        i += 1
                        continue

                    # Check for section header
                    if t.upper() in _MONSTER_SECTION_HEADERS:
                        current_section = t.lower()
                        sections.setdefault(current_section, [])
                        i += 1
                        continue

                    # Check if this looks like the next monster's name
                    # (non-uppercase, non-empty, followed by a stat table)
                    if (i + 1 < len(elements)
                            and isinstance(elements[i + 1], DocxTable)
                            and _is_monster_stat_table(elements[i + 1])):
                        break

                    sections.setdefault(current_section, []).append(t)

                i += 1

            # Build monster dict
            monster = _build_monster_dict(monster_name, stats, sections)
            monsters.append(monster)
            continue

        i += 1

    return monsters


def _build_monster_dict(name: str, stats: dict, sections: dict) -> dict:
    img_parts = sections.get("image prompt", [])
    image_prompt = img_parts[0] if img_parts else ""

    return {
        "name": name,
        "monster_type": stats.get("Monster Type", ""),
        "challenge_rating": stats.get("Challenge Rating", ""),
        "armor_class": stats.get("Armor Class", ""),
        "hit_points": stats.get("Hit Points", ""),
        "speed": stats.get("Speed", ""),
        "ability_scores": stats.get("Ability Scores", ""),
        "saving_throws": stats.get("Saving Throws", ""),
        "skills": stats.get("Skills", ""),
        "damage_resistances": stats.get("Damage Resistances", ""),
        "damage_immunities": stats.get("Damage Immunities", ""),
        "condition_immunities": stats.get("Condition Immunities", ""),
        "senses": stats.get("Senses", ""),
        "languages": stats.get("Languages", ""),
        "dungeon_type": stats.get("Dungeon Type", ""),
        "faction": stats.get("Faction", ""),
        "traits": "\n".join(sections.get("traits", [])),
        "actions": "\n".join(sections.get("actions", [])),
        "reactions": "\n".join(sections.get("reactions", [])),
        "lore": "\n".join(sections.get("lore", [])),
        "ecology": "\n".join(sections.get("ecology", [])),
        "encounter_hooks": "\n".join(sections.get("encounter hooks", [])),
        "image_prompt": image_prompt,
    }


def _monster_to_text(monster: dict) -> str:
    lines = [f"Name: {monster['name']}"]
    skip = {"name", "image_prompt"}
    for key, value in monster.items():
        if key in skip or not value:
            continue
        label = key.replace("_", " ").title()
        lines.append(f"{label}: {value}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Artifact parser
# ---------------------------------------------------------------------------

_ARTIFACT_META_KEYS = {
    "Artifact Type", "Item Subtype", "Rarity", "Attunement",
    "Attunement Req.", "Creator / Origin", "Current Location",
}

_ARTIFACT_SECTIONS = {
    "PHYSICAL DESCRIPTION", "PROPERTIES", "POWERS & ABILITIES",
    "SPELLS", "LORE & HISTORY", "SENTIENCE", "CONSEQUENCES",
    "RUMORS & LEGENDS", "ENCOUNTER HOOKS", "RAG / KNOWLEDGE BASE",
}


def _is_artifact_meta_table(table: DocxTable) -> bool:
    first_col = {row.cells[0].text.strip() for row in table.rows if row.cells}
    return bool(first_col & {"Artifact Type", "Rarity", "Attunement"})


def _is_artifact_title_table(table: DocxTable) -> bool:
    """Title tables are 1-row tables with the artifact name."""
    if len(table.rows) != 1:
        return False
    text = table.rows[0].cells[0].text.strip()
    # Title block has name on first line, type info on second
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    return len(lines) >= 2 and "·" in lines[1]


def extract_artifacts(docx_path: Path) -> list[dict]:
    doc = Document(docx_path)
    elements = list(_iter_body(doc))
    artifacts = []
    i = 0

    while i < len(elements):
        el = elements[i]

        # Look for title tables (1 row with "Name\nType · Subtype · Rarity")
        if isinstance(el, DocxTable) and _is_artifact_title_table(el):
            cell_text = el.rows[0].cells[0].text.strip()
            lines = [l.strip() for l in cell_text.split("\n") if l.strip()]
            artifact_name = lines[0]
            i += 1

            # Next should be the metadata table
            meta = {}
            if i < len(elements) and isinstance(elements[i], DocxTable):
                if _is_artifact_meta_table(elements[i]):
                    meta = _table_to_dict(elements[i])
                    i += 1

            # Collect sections (1-row header tables followed by paragraphs)
            sections = {}
            current_section = None

            while i < len(elements):
                el2 = elements[i]

                if isinstance(el2, DocxTable):
                    # Check if next artifact title
                    if _is_artifact_title_table(el2):
                        break

                    # Check if section header (1-row table)
                    header = _table_cell_text(el2).strip()
                    if header in _ARTIFACT_SECTIONS:
                        current_section = header.lower()
                        sections.setdefault(current_section, [])
                        i += 1
                        continue

                    # Could be metadata table for this artifact
                    if _is_artifact_meta_table(el2) and not meta:
                        meta = _table_to_dict(el2)

                    i += 1
                    continue

                if isinstance(el2, DocxParagraph):
                    t = el2.text.strip()
                    if t and current_section:
                        sections.setdefault(current_section, []).append(t)
                    i += 1
                    continue

                i += 1

            artifact = _build_artifact_dict(artifact_name, meta, sections)
            artifacts.append(artifact)
            continue

        i += 1

    return artifacts


def _build_artifact_dict(name: str, meta: dict, sections: dict) -> dict:
    # Parse rumors and encounter hooks as lists
    rumors_text = "\n".join(sections.get("rumors & legends", []))
    rumors = [r.strip() for r in rumors_text.split("\n") if r.strip()] if rumors_text else []

    hooks_text = "\n".join(sections.get("encounter hooks", []))
    hooks = [h.strip() for h in hooks_text.split("\n") if h.strip()] if hooks_text else []

    rag_parts = sections.get("rag / knowledge base", [])
    image_prompt = ""
    tags = ""
    for part in rag_parts:
        if part.lower().startswith("image prompt:"):
            image_prompt = part.split(":", 1)[1].strip()
        elif part.lower().startswith("tags:"):
            tags = part.split(":", 1)[1].strip()

    return {
        "name": name,
        "artifact_type": meta.get("Artifact Type", ""),
        "item_subtype": meta.get("Item Subtype", ""),
        "rarity": meta.get("Rarity", ""),
        "attunement": meta.get("Attunement", ""),
        "attunement_req": meta.get("Attunement Req.", ""),
        "creator_origin": meta.get("Creator / Origin", ""),
        "current_location": meta.get("Current Location", ""),
        "physical_description": "\n".join(sections.get("physical description", [])),
        "properties": "\n".join(sections.get("properties", [])),
        "powers_and_abilities": "\n".join(sections.get("powers & abilities", [])),
        "spells": "\n".join(sections.get("spells", [])),
        "lore_and_history": "\n".join(sections.get("lore & history", [])),
        "sentience": "\n".join(sections.get("sentience", [])),
        "consequences": "\n".join(sections.get("consequences", [])),
        "rumors": rumors,
        "encounter_hooks": hooks,
        "image_prompt": image_prompt,
        "tags": tags,
    }


def _artifact_to_text(artifact: dict) -> str:
    lines = [f"Name: {artifact['name']}"]
    skip = {"name", "image_prompt", "tags", "rumors", "encounter_hooks"}
    for key, value in artifact.items():
        if key in skip or not value:
            continue
        label = key.replace("_", " ").title()
        lines.append(f"{label}: {value}")
    if artifact.get("rumors"):
        lines.append("Rumors: " + " | ".join(artifact["rumors"]))
    if artifact.get("encounter_hooks"):
        lines.append("Encounter Hooks: " + " | ".join(artifact["encounter_hooks"]))
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Unified extraction dispatch
# ---------------------------------------------------------------------------

EXTRACTORS = {
    "weapon":   extract_weapons,
    "npc":      extract_npcs,
    "location": extract_locations,
    "monster":  extract_monsters,
    "artifact": extract_artifacts,
}

TEXT_CONVERTERS = {
    "weapon":   _weapon_to_text,
    "npc":      _npc_to_text,
    "location": _location_to_text,
    "monster":  _monster_to_text,
    "artifact": _artifact_to_text,
}


def extract_category(category: str) -> list[dict]:
    """Extract all items from a category's docx files."""
    dir_name = CATEGORY_DIRS[category]
    data_path = DATA_DIR / dir_name
    extractor = EXTRACTORS[category]

    docx_files = sorted(data_path.glob("*.docx"))
    if not docx_files:
        console.print(f"[yellow]No .docx files found in {data_path}[/yellow]")
        return []

    all_items = []
    for path in docx_files:
        items = extractor(path)
        all_items.extend(items)
        console.print(f"  {path.name}: [cyan]{len(items)}[/cyan] items")

    return all_items


# ---------------------------------------------------------------------------
# LlamaIndex + ChromaDB indexing
# ---------------------------------------------------------------------------

def build_index(items: list[dict], category: str) -> int:
    """Embed items and store them in ChromaDB with category metadata."""
    CHROMA_PATH.mkdir(parents=True, exist_ok=True)

    chroma_client = chromadb.PersistentClient(path=str(CHROMA_PATH))
    chroma_collection = chroma_client.get_or_create_collection(COLLECTION_NAME)

    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    Settings.embed_model = OpenAIEmbedding(model="text-embedding-3-small")
    Settings.llm = None  # ingestion only

    to_text = TEXT_CONVERTERS[category]

    documents = []
    for item in items:
        doc = LlamaDocument(
            text=to_text(item),
            metadata={
                "category": category,
                "name": item.get("name", "Unknown"),
                "rarity": item.get("rarity", ""),
                "type": item.get("item_type", item.get("artifact_type",
                        item.get("location_type", item.get("monster_type", "")))),
            },
        )
        documents.append(doc)

    if documents:
        VectorStoreIndex.from_documents(
            documents,
            storage_context=storage_context,
            show_progress=True,
        )

    return len(documents)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Ingest D&D content into ChromaDB")
    parser.add_argument(
        "--category",
        choices=list(CATEGORY_DIRS.keys()),
        default=None,
        help="Ingest only this category (default: all)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Parse and count items without embedding",
    )
    args = parser.parse_args()

    categories = [args.category] if args.category else list(CATEGORY_DIRS.keys())

    console.rule("[bold cyan]D&D Knowledge Base Ingestion[/bold cyan]")

    summary = Table("Category", "Files", "Items", show_header=True, header_style="bold magenta")

    for category in categories:
        console.print(f"\n[bold]Processing [cyan]{category}[/cyan]…[/bold]")
        items = extract_category(category)
        dir_name = CATEGORY_DIRS[category]
        file_count = len(list((DATA_DIR / dir_name).glob("*.docx")))
        summary.add_row(category, str(file_count), str(len(items)))

        if not items:
            console.print(f"[yellow]No items extracted for {category}[/yellow]")
            continue

        if args.dry_run:
            console.print(f"  [dim](dry run — skipping embedding)[/dim]")
            continue

        console.print(f"  Embedding [cyan]{len(items)}[/cyan] items…")
        indexed = build_index(items, category)
        console.print(f"  [green]✓[/green] {indexed} items indexed")

    console.print()
    console.print(summary)
    console.rule("[bold green]Done[/bold green]")

    if not args.dry_run:
        console.print(f"[dim]Vector store persisted at: {CHROMA_PATH.resolve()}[/dim]")


if __name__ == "__main__":
    main()
